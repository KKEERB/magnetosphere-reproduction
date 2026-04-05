"""
车企配件采购计划半自动生成系统
Phase 5 & 6: Excel 和 Word 输出模块

功能：
- 生成采购计划 Excel（多 Sheet）
- 生成 Word 说明文档
- 格式化输出
"""

import pandas as pd
import numpy as np
from pathlib import Path
from datetime import datetime
import logging

# 尝试导入 python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("未安装 python-docx，Word 输出功能将不可用")

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/export_module.log', encoding='utf-8', mode='w'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class ExcelExporter:
    """采购计划 Excel 导出器"""

    def __init__(self, procurement_data: pd.DataFrame):
        """
        初始化导出器

        Args:
            procurement_data: 采购计算结果 DataFrame
        """
        self.data = procurement_data
        self.writer = None

    def export(self, output_path: str) -> str:
        """
        导出采购计划 Excel

        Args:
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        logger.info(f"开始导出 Excel 到：{output_path}")

        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            self.writer = writer

            # Sheet 1: 采购计划总表
            self._export_summary_sheet()

            # Sheet 2: 高优先级配件明细
            self._export_high_priority_sheet()

            # Sheet 3: 紧急采购清单
            self._export_urgent_list()

            # Sheet 4: ABC 分类统计表
            self._export_abc_summary()

            # Sheet 5: 模型参数表
            self._export_model_params()

        logger.info(f"Excel 导出完成：{output_path}")
        return output_path

    def _export_summary_sheet(self):
        """导出采购计划总表"""
        logger.info("导出 Sheet 1: 采购计划总表")

        # 选择需要的列
        cols_to_export = [
            '配件号', '名称', '单位', 'ABC 分类', '采购优先级',
            '月均需求', '日均需求', '最大储备量', '最小储备量',
            '建议采购量', '预测模型', '置信度'
        ]

        # 过滤存在的列
        available_cols = [c for c in cols_to_export if c in self.data.columns]
        df = self.data[available_cols].copy()

        # 只导出需要采购的配件（建议采购量 > 0）或高优先级配件
        # 为了完整性，导出所有配件
        df.to_excel(self.writer, sheet_name='采购计划总表', index=False)

    def _export_high_priority_sheet(self):
        """导出高优先级配件明细"""
        logger.info("导出 Sheet 2: 高优先级配件明细")

        # 过滤 A 类和需要采购的配件
        high_priority_mask = (
            (self.data['ABC 分类'] == 'A') |
            (self.data['建议采购量'] > 0)
        )

        high_priority_df = self.data[high_priority_mask].copy()

        # 按采购优先级和采购量排序
        if '采购优先级' in high_priority_df.columns:
            priority_order = {'高': 0, '中': 1, '低': 2}
            high_priority_df['priority_sort'] = high_priority_df['采购优先级'].map(priority_order)
            high_priority_df = high_priority_df.sort_values(['priority_sort', '建议采购量'], ascending=[True, False])
            high_priority_df = high_priority_df.drop(columns=['priority_sort'])

        high_priority_df.to_excel(self.writer, sheet_name='高优先级明细', index=False)

    def _export_urgent_list(self):
        """导出紧急采购清单"""
        logger.info("导出 Sheet 3: 紧急采购清单")

        # 紧急采购：采购优先级为高 且 建议采购量 > 0
        urgent_mask = (
            (self.data['采购优先级'] == '高') &
            (self.data['建议采购量'] > 0)
        )

        urgent_df = self.data[urgent_mask].copy()

        if len(urgent_df) == 0:
            # 创建空表头
            urgent_df = pd.DataFrame(columns=self.data.columns)

        urgent_df.to_excel(self.writer, sheet_name='紧急采购清单', index=False)

    def _export_abc_summary(self):
        """导出 ABC 分类统计表"""
        logger.info("导出 Sheet 4: ABC 分类统计表")

        # 按 ABC 分类汇总
        abc_summary = self.data.groupby('ABC 分类').agg({
            '配件号': 'count',
            '月均需求': 'sum',
            '建议采购量': 'sum',
        }).reset_index()

        abc_summary.columns = ['ABC 分类', '配件数量', '月均需求总量', '建议采购总量']
        abc_summary.to_excel(self.writer, sheet_name='ABC 分类统计', index=False)

    def _export_model_params(self):
        """导出模型参数表"""
        logger.info("导出 Sheet 5: 模型参数表")

        # 模型使用情况统计
        if '预测模型' in self.data.columns:
            model_summary = self.data.groupby('预测模型').agg({
                '配件号': 'count',
                '置信度': 'mean',
                '建议采购量': 'sum'
            }).reset_index()
            model_summary.columns = ['预测模型', '配件数量', '平均置信度', '建议采购总量']
            model_summary.to_excel(self.writer, sheet_name='模型参数表', index=False)


class WordExporter:
    """采购计划 Word 说明文档导出器"""

    def __init__(self, procurement_data: pd.DataFrame, stats: dict):
        """
        初始化导出器

        Args:
            procurement_data: 采购计算结果 DataFrame
            stats: 统计摘要字典
        """
        self.data = procurement_data
        self.stats = stats
        self.doc = None

    def export(self, output_path: str) -> str:
        """
        导出 Word 说明文档

        Args:
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        if not HAS_DOCX:
            logger.error("python-docx 未安装，无法生成 Word 文档")
            # 创建一个文本文件作为替代
            self._export_text_report(output_path.replace('.docx', '.txt'))
            return output_path.replace('.docx', '.txt')

        logger.info(f"开始导出 Word 文档到：{output_path}")

        self.doc = Document()

        # 1. 标题
        self._add_title()

        # 2. 执行摘要
        self._add_executive_summary()

        # 3. 数据说明
        self._add_data_description()

        # 4. 预测模型说明
        self._add_model_description()

        # 5. ABC 分类说明
        self._add_abc_description()

        # 6. 风险提示
        self._add_risk_warnings()

        # 7. 保存文档
        self.doc.save(output_path)
        logger.info(f"Word 文档导出完成：{output_path}")

        return output_path

    def _add_title(self):
        """添加标题"""
        title = self.doc.add_heading('配件采购计划说明', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 日期
        date_str = datetime.now().strftime('%Y年%m月%d日')
        date_para = self.doc.add_paragraph(f'生成日期：{date_str}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()  # 空行

    def _add_executive_summary(self):
        """添加执行摘要"""
        self.doc.add_heading('1. 执行摘要', level=1)

        # 创建摘要表格
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # 填充数据
        cells = [
            ('总配件数', f"{self.stats.get('total_items', 0):,} 件"),
            ('需要采购', f"{self.stats.get('items_to_procure', 0):,} 件"),
            ('总采购量', f"{self.stats.get('total_procurement_qty', 0):,} 件"),
            ('高优先级', f"{self.stats.get('high_priority_count', 0):,} 件"),
        ]

        for i, (label, value) in enumerate(cells):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

    def _add_data_description(self):
        """添加数据说明"""
        self.doc.add_heading('2. 数据说明', level=1)

        self.doc.add_paragraph('数据来源：')
        self.doc.add_paragraph('• 历史采购数据：2016 年 1 月 - 2026 年 2 月（约 10 年）', style='List Bullet')
        self.doc.add_paragraph('• 配件品类：约 55,895 种', style='List Bullet')
        self.doc.add_paragraph('• 数据质量：已清洗和标准化处理', style='List Bullet')

        self.doc.add_paragraph()

    def _add_model_description(self):
        """添加预测模型说明"""
        self.doc.add_heading('3. 预测模型说明', level=1)

        self.doc.add_paragraph('本系统采用数据驱动的预测方法，根据配件需求特征自动选择最佳模型：')

        self.doc.add_paragraph('• 移动平均法 (SMA)', style='List Bullet')
        self.doc.add_paragraph('  适用于需求稳定的配件，计算过去 N 个月的平均值', style='Intense Quote')

        self.doc.add_paragraph('• 指数平滑法', style='List Bullet')
        self.doc.add_paragraph('  适用于有趋势但无明显季节性的配件', style='Intense Quote')

        self.doc.add_paragraph('• Croston 法', style='List Bullet')
        self.doc.add_paragraph('  适用于间歇性需求（需求不连续、偶尔发生）的配件', style='Intense Quote')

        self.doc.add_paragraph()

    def _add_abc_description(self):
        """添加 ABC 分类说明"""
        self.doc.add_heading('4. ABC 分类说明', level=1)

        self.doc.add_paragraph('基于 12 个月出库频次 (T) 的分类规则：')

        abc_rules = [
            ('A 类', 'T ≥ 120 次', '重点管理、精确预测'),
            ('B 类', '36 ≤ T < 120', '常规管理'),
            ('C 类', '12 ≤ T < 36', '简化管理'),
            ('D 类', '1 ≤ T < 12', '简化管理'),
            ('E 类', 'T = 0', '无需求、暂停储备'),
        ]

        table = self.doc.add_table(rows=len(abc_rules) + 1, cols=3)
        table.style = 'Table Grid'

        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '分类'
        header_cells[1].text = '条件'
        header_cells[2].text = '管理策略'

        # 数据行
        for i, (grade, condition, strategy) in enumerate(abc_rules):
            row = table.rows[i + 1]
            row.cells[0].text = grade
            row.cells[1].text = condition
            row.cells[2].text = strategy

        self.doc.add_paragraph()

    def _add_risk_warnings(self):
        """添加风险提示"""
        self.doc.add_heading('5. 风险提示', level=1)

        self.doc.add_paragraph('请注意以下事项：')

        self.doc.add_paragraph('• 预测结果基于历史数据，实际情况可能因市场变化而有所不同', style='List Bullet')
        self.doc.add_paragraph('• 无历史需求记录的配件（E 类）未纳入采购建议', style='List Bullet')
        self.doc.add_paragraph('• 建议结合实际情况和专业知识进行人工审核', style='List Bullet')
        self.doc.add_paragraph('• 紧急采购清单中的配件需要优先处理', style='List Bullet')

        self.doc.add_paragraph()
        self.doc.add_paragraph('---')
        self.doc.add_paragraph('本报告由车企配件采购计划半自动生成系统生成')

    def _export_text_report(self, output_path: str):
        """导出文本报告（当 python-docx 不可用时的备选方案）"""
        logger.info(f"导出文本报告到：{output_path}")

        content = []
        content.append("=" * 60)
        content.append("配件采购计划说明")
        content.append("=" * 60)
        content.append(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
        content.append("")

        content.append("1. 执行摘要")
        content.append("-" * 40)
        content.append(f"总配件数：{self.stats.get('total_items', 0):,} 件")
        content.append(f"需要采购：{self.stats.get('items_to_procure', 0):,} 件")
        content.append(f"总采购量：{self.stats.get('total_procurement_qty', 0):,} 件")
        content.append("")

        content.append("2. 数据说明")
        content.append("-" * 40)
        content.append("数据来源：历史采购数据 2016 年 1 月 - 2026 年 2 月")
        content.append("配件品类：约 55,895 种")
        content.append("")

        content.append("3. 预测模型")
        content.append("-" * 40)
        content.append("- 移动平均法 (SMA): 适用于需求稳定的配件")
        content.append("- 指数平滑法：适用于有趋势的需求")
        content.append("- Croston 法：适用于间歇性需求")
        content.append("")

        content.append("4. ABC 分类规则")
        content.append("-" * 40)
        content.append("A 类：T ≥ 120 次 - 重点管理")
        content.append("B 类：36 ≤ T < 120 - 常规管理")
        content.append("C 类：12 ≤ T < 36 - 简化管理")
        content.append("D 类：1 ≤ T < 12 - 简化管理")
        content.append("E 类：T = 0 - 无需求")
        content.append("")

        content.append("5. 风险提示")
        content.append("-" * 40)
        content.append("- 预测结果基于历史数据，实际情况可能有所不同")
        content.append("- 建议结合实际情况进行人工审核")
        content.append("")

        content.append("=" * 60)
        content.append("车企配件采购计划半自动生成系统")
        content.append("=" * 60)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))

        logger.info(f"文本报告导出完成：{output_path}")


def main():
    """主函数 - 测试输出模块"""

    # 配置路径
    input_dir = Path("C:/Users/Lenovo/配件采购系统/data/output")
    output_dir = Path("C:/Users/Lenovo/配件采购系统/data/output")

    logger.info("=" * 60)
    logger.info("车企配件采购系统 - Excel 和 Word 输出模块")
    logger.info("=" * 60)

    # 1. 加载采购计算结果
    print("\n[1/4] 加载采购计算结果...")
    procurement_files = sorted(input_dir.glob("procurement_calc_*.xlsx"))
    if not procurement_files:
        raise FileNotFoundError("未找到采购计算结果，请先运行 Phase 4")

    latest_file = procurement_files[-1]
    logger.info(f"加载文件：{latest_file.name}")
    data = pd.read_excel(latest_file, engine='openpyxl')
    logger.info(f"数据行数：{len(data)}")

    # 2. 计算统计摘要
    print("\n[2/4] 计算统计摘要...")
    stats = {
        'total_items': len(data),
        'items_to_procure': int((data['建议采购量'] > 0).sum()),
        'total_procurement_qty': int(data['建议采购量'].sum()),
        'high_priority_count': int((data['采购优先级'] == '高').sum()),
        'medium_priority_count': int((data['采购优先级'] == '中').sum()),
        'low_priority_count': int((data['采购优先级'] == '低').sum()),
    }

    # ABC 分类采购量统计
    abc_procurement = data.groupby('ABC 分类')['建议采购量'].sum().to_dict()
    stats['abc_procurement'] = abc_procurement

    print(f"  总配件数：{stats['total_items']:,} 件")
    print(f"  需要采购：{stats['items_to_procure']:,} 件")
    print(f"  总采购量：{stats['total_procurement_qty']:,} 件")

    # 3. 导出 Excel
    print("\n[3/4] 导出 Excel 采购计划...")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    excel_exporter = ExcelExporter(data)
    excel_file = output_dir / f"采购计划_{timestamp}.xlsx"
    excel_exporter.export(str(excel_file))
    print(f"  Excel 已保存：{excel_file.name}")

    # 4. 导出 Word
    print("\n[4/4] 导出 Word 说明文档...")
    word_exporter = WordExporter(data, stats)
    word_file = output_dir / f"采购计划说明_{timestamp}.docx"
    word_file_path = word_exporter.export(str(word_file))
    print(f"  Word 文档已保存：{Path(word_file_path).name}")

    logger.info("=" * 60)
    logger.info("Phase 5 & 6 完成!")
    logger.info("=" * 60)

    print("\n" + "=" * 60)
    print("所有阶段完成！输出文件：")
    print(f"  - Excel: {excel_file.name}")
    print(f"  - Word:  {Path(word_file_path).name}")
    print("=" * 60)

    return excel_file, word_file_path


if __name__ == "__main__":
    main()
